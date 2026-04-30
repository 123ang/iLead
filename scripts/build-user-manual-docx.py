from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re

ROOT = Path(__file__).resolve().parents[1]
manual_path = ROOT / 'USER_MANUAL.md'
screens_dir = ROOT / 'docs' / 'user-manual-screenshots'
out_path = ROOT / 'docs' / 'iLead_User_Manual_with_Screenshots.docx'
out_path.parent.mkdir(parents=True, exist_ok=True)

SCREENSHOT_BY_HEADING = {
    '## 2) Sign in': [('01-login.png', 'Login screen')],
    '### Dashboard': [('02-dashboard.png', 'Executive dashboard with yearly graphs')],
    '### Campaigns': [('03-campaigns.png', 'Campaign management list')],
    '### Leads': [('04-leads.png', 'Lead management list')],
    '### Applications (Upload)': [
        ('05-application-upload.png', 'Application upload screen'),
        ('06-application-upload-mapping.png', 'Column mapping after choosing a CSV file'),
    ],
    '### Follow-ups': [('07-follow-ups.png', 'Follow-up overdue queue')],
    '### Duplicates': [('08-duplicates.png', 'Duplicate leads review')],
    '### Reports': [('09-reports.png', 'Management reports and CSV export')],
    '### Master Data': [('10-master-data.png', 'Master data administration')],
    '### Users': [('11-users.png', 'User management')],
    '### Settings': [('12-settings.png', 'System settings')],
    '### Audit Logs': [('13-audit-logs.png', 'Audit logs')],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_run_font(run, name='Aptos', size=None, color=None, bold=None):
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if bold is not None:
        run.bold = bold


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(80, 90, 105)


def add_screenshot(doc, filename, caption):
    img = screens_dir / filename
    if not img.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Landscape A4 minus margins; keep screenshots readable.
    run.add_picture(str(img), width=Inches(10.2))
    add_caption(doc, f'Screenshot: {caption}')


def add_screens_for_heading(doc, heading_line):
    for filename, caption in SCREENSHOT_BY_HEADING.get(heading_line, []):
        add_screenshot(doc, filename, caption)


def add_inline_markdown(paragraph, text):
    # Simple bold/code handling good enough for this manual.
    parts = re.split(r'(`[^`]+`|\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11.69)
section.page_height = Inches(8.27)
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.45)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)

styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal'].font.size = Pt(10)
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    styles[style_name].font.name = 'Aptos Display'
    styles[style_name].font.color.rgb = RGBColor(7, 27, 58)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('iLead')
set_run_font(r, 'Aptos Display', 30, (7, 27, 58), True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('First-Time User Manual')
set_run_font(r, 'Aptos Display', 20, (201, 162, 39), True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MVP / UAT Version with Screenshots')
set_run_font(r, 'Aptos', 12, (70, 80, 95), False)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Generated: 2026-04-29')
set_run_font(r, 'Aptos', 10, (70, 80, 95), False)

doc.add_paragraph()
add_screenshot(doc, '02-dashboard.png', 'Executive dashboard overview')
doc.add_page_break()

# Quick screenshot index table
h = doc.add_heading('Screenshot Index', level=1)
intro = doc.add_paragraph('This document includes screenshots captured from the local MVP/UAT environment using the seeded demo admin account.')

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, text in enumerate(['No.', 'Screen', 'Purpose']):
    hdr[i].text = text
    set_cell_shading(hdr[i], '071B3A')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

index_rows = [
    ('1', 'Login', 'Sign in to iLead'),
    ('2', 'Dashboard', 'Executive KPIs and yearly graphs'),
    ('3', 'Campaigns', 'Campaign list and management'),
    ('4', 'Leads', 'Lead list and management'),
    ('5', 'Applications', 'Upload CSV/XLSX application files'),
    ('6', 'Column Mapping', 'Map uploaded file columns to iLead fields'),
    ('7', 'Follow-ups', 'Overdue queue and follow-up actions'),
    ('8', 'Duplicates', 'Duplicate lead review'),
    ('9', 'Reports', 'Management reports and CSV export'),
    ('10', 'Master Data', 'Admin reference data'),
    ('11', 'Users', 'User account management'),
    ('12', 'Settings', 'System settings'),
    ('13', 'Audit Logs', 'Action and export audit trail'),
]
for row in index_rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text

doc.add_page_break()

in_code = False
code_lines = []

for raw in manual_path.read_text(encoding='utf-8').splitlines():
    line = raw.rstrip('\n')

    if line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            for c in code_lines:
                run = p.add_run(c + '\n')
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
            code_lines = []
        continue

    if in_code:
        code_lines.append(line)
        continue

    if not line.strip():
        doc.add_paragraph()
        continue

    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
        continue
    if line.startswith('## '):
        full_heading = line.strip()
        doc.add_heading(line[3:].strip(), level=1)
        add_screens_for_heading(doc, full_heading)
        continue
    if line.startswith('### '):
        full_heading = line.strip()
        doc.add_heading(line[4:].strip(), level=2)
        add_screens_for_heading(doc, full_heading)
        continue

    if line.startswith('>'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line.lstrip('> ').strip())
        run.italic = True
        run.font.color.rgb = RGBColor(80, 90, 105)
        continue

    if re.match(r'^\d+\.\s+', line):
        p = doc.add_paragraph(style='List Number')
        add_inline_markdown(p, re.sub(r'^\d+\.\s+', '', line))
        continue

    if line.strip().startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        add_inline_markdown(p, line.strip()[2:])
        continue

    p = doc.add_paragraph()
    add_inline_markdown(p, line)

# Add a final MVP/UAT note.
doc.add_page_break()
doc.add_heading('MVP / UAT Notes', level=1)
for bullet in [
    'This manual is intended for MVP/UAT use with the seeded local environment.',
    'Some production hardening items are tracked separately in ISSUES_FOUND.md.',
    'For UAT, focus on validating navigation, campaign/lead workflows, application upload, matching, follow-ups, duplicates, reports, users, settings, and audit logs.',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bullet)

doc.save(out_path)
print(out_path)
